from __future__ import annotations

import csv
from pathlib import Path
from zipfile import BadZipFile

from app.parsers.base import ParsedSection
from app.parsers.markdown import parse_markdown_sections


class DoclingParseError(RuntimeError):
    pass


class DoclingParser:
    """Local Docling parser for simple office documents."""

    def parse_sections(self, file_path: str, file_name: str) -> list[ParsedSection]:
        markdown = self.convert_to_markdown(file_path, file_name)
        return parse_markdown_sections(markdown)

    def convert_to_markdown(self, file_path: str, file_name: str) -> str:
        suffix = Path(file_name).suffix.lower()
        if suffix not in {".docx", ".xlsx", ".csv"}:
            raise DoclingParseError(f"Local Docling parser supports only docx, xlsx, and csv; got {suffix}.")

        try:
            return _convert_with_docling_backend(file_path, file_name)
        except Exception:
            return _convert_office_to_markdown(file_path, file_name)


def _convert_with_docling_backend(file_path: str, file_name: str) -> str:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".docx":
        from docling.backend.msword_backend import MsWordDocumentBackend
        from docling.datamodel.base_models import InputFormat

        return _run_docling_backend(file_path, file_name, InputFormat.DOCX, MsWordDocumentBackend)
    if suffix == ".xlsx":
        from docling.backend.msexcel_backend import MsExcelDocumentBackend
        from docling.datamodel.base_models import InputFormat

        return _run_docling_backend(file_path, file_name, InputFormat.XLSX, MsExcelDocumentBackend)
    if suffix == ".csv":
        return _csv_to_markdown(file_path)
    raise DoclingParseError(f"Unsupported office file type: {suffix}.")


def _run_docling_backend(file_path: str, file_name: str, input_format, backend_cls) -> str:
    from docling.datamodel.document import InputDocument

    path = Path(file_path)
    input_document = InputDocument(path, input_format, backend_cls, filename=file_name)
    backend = backend_cls(input_document, path)
    document = backend.convert()
    markdown = document.export_to_markdown()
    if not markdown.strip():
        raise DoclingParseError(f"Local Docling conversion returned empty markdown for {file_name}.")
    return markdown


def _convert_office_to_markdown(file_path: str, file_name: str) -> str:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".docx":
        return _docx_to_markdown(file_path)
    if suffix == ".xlsx":
        return _xlsx_to_markdown(file_path)
    if suffix == ".csv":
        return _csv_to_markdown(file_path)
    raise DoclingParseError(f"Unsupported office file type: {suffix}.")


def _docx_to_markdown(file_path: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DoclingParseError("python-docx is not installed.") from exc

    try:
        document = Document(file_path)
    except (BadZipFile, Exception) as exc:
        raise DoclingParseError(f"Failed to read docx: {exc}") from exc

    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        if style_name.startswith("heading"):
            level = _heading_level(style_name)
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text)

    for table in document.tables:
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        markdown_table = _rows_to_markdown_table(rows)
        if markdown_table:
            lines.append(markdown_table)

    return "\n\n".join(lines)


def _xlsx_to_markdown(file_path: str) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise DoclingParseError("openpyxl is not installed.") from exc

    workbook = load_workbook(file_path, read_only=True, data_only=True)
    blocks: list[str] = []
    for worksheet in workbook.worksheets:
        rows = [
            ["" if value is None else str(value) for value in row]
            for row in worksheet.iter_rows(values_only=True)
        ]
        rows = _trim_empty_rows(rows)
        if not rows:
            continue
        table = _rows_to_markdown_table(rows)
        if table:
            blocks.append(f"# {worksheet.title}\n\n{table}")
    workbook.close()
    return "\n\n".join(blocks)


def _csv_to_markdown(file_path: str) -> str:
    rows: list[list[str]] = []
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            with open(file_path, newline="", encoding=encoding) as f:
                rows = [row for row in csv.reader(f)]
            break
        except UnicodeDecodeError:
            continue
    rows = _trim_empty_rows(rows)
    return _rows_to_markdown_table(rows)


def _heading_level(style_name: str) -> int:
    parts = style_name.split()
    if parts and parts[-1].isdigit():
        return min(max(int(parts[-1]), 1), 3)
    return 1


def _trim_empty_rows(rows: list[list[str]]) -> list[list[str]]:
    return [row for row in rows if any(str(cell).strip() for cell in row)]


def _rows_to_markdown_table(rows: list[list[str]]) -> str:
    rows = _trim_empty_rows(rows)
    if not rows:
        return ""

    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = [_escape_table_cell(cell) for cell in normalized[0]]
    separator = ["---"] * width
    body = [[_escape_table_cell(cell) for cell in row] for row in normalized[1:]]
    lines = [
        f"| {' | '.join(header)} |",
        f"| {' | '.join(separator)} |",
    ]
    lines.extend(f"| {' | '.join(row)} |" for row in body)
    return "\n".join(lines)


def _escape_table_cell(value: object) -> str:
    return str(value).replace("|", "\\|").replace("\n", " ").strip()
